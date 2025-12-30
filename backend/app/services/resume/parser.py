"""
Resume Parser - Extract text and sections from various file formats
"""

from pathlib import Path
from typing import Dict, Optional, List
import logging
import re

logger = logging.getLogger(__name__)


class ResumeParser:
    """
    Parse resumes from various file formats.
    Supports: PDF, DOCX, DOC, TXT
    """

    # Common section headers for detection
    SECTION_PATTERNS = {
        "contact": r"(?i)(contact\s*info|contact\s*details|personal\s*info)",
        "summary": r"(?i)(summary|profile|objective|about\s*me|professional\s*summary)",
        "experience": r"(?i)(experience|work\s*history|employment|professional\s*experience|career\s*history)",
        "education": r"(?i)(education|academic|qualifications|degrees)",
        "skills": r"(?i)(skills|technical\s*skills|core\s*competencies|expertise|technologies)",
        "projects": r"(?i)(projects|portfolio|key\s*projects)",
        "certifications": r"(?i)(certifications|certificates|licenses|credentials)",
        "awards": r"(?i)(awards|honors|achievements|recognition)",
        "publications": r"(?i)(publications|papers|research)",
        "languages": r"(?i)(languages|language\s*skills)",
        "interests": r"(?i)(interests|hobbies|activities)"
    }

    async def parse(self, file_path: Path) -> Dict:
        """
        Parse a resume file and extract content.

        Args:
            file_path: Path to the resume file

        Returns:
            Dict containing:
                - text: Full extracted text
                - sections: Detected sections with content
                - metadata: File metadata
        """
        file_ext = file_path.suffix.lower()

        # Extract text based on file type
        if file_ext == ".pdf":
            text = await self._parse_pdf(file_path)
        elif file_ext in [".docx", ".doc"]:
            text = await self._parse_docx(file_path)
        elif file_ext == ".txt":
            text = await self._parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

        # Clean and normalize text
        text = self._clean_text(text)

        # Detect sections
        sections = self._detect_sections(text)

        # Extract contact info
        contact_info = self._extract_contact_info(text)

        return {
            "text": text,
            "sections": sections,
            "contact_info": contact_info,
            "metadata": {
                "filename": file_path.name,
                "file_type": file_ext,
                "word_count": len(text.split()),
                "char_count": len(text)
            }
        }

    async def _parse_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            import pdfplumber

            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

            return "\n".join(text_parts)

        except ImportError:
            # Fallback to PyPDF2
            try:
                from PyPDF2 import PdfReader

                text_parts = []
                reader = PdfReader(file_path)
                for page in reader.pages:
                    text_parts.append(page.extract_text())

                return "\n".join(text_parts)

            except ImportError:
                logger.error("No PDF library available. Install pdfplumber or PyPDF2")
                raise ImportError("PDF parsing requires pdfplumber or PyPDF2")

    async def _parse_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            from docx import Document

            doc = Document(file_path)
            text_parts = []

            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)

            return "\n".join(text_parts)

        except ImportError:
            logger.error("python-docx not installed")
            raise ImportError("DOCX parsing requires python-docx")

    async def _parse_txt(self, file_path: Path) -> str:
        """Read plain text file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)

        # Remove non-printable characters
        text = re.sub(r'[^\x20-\x7E\n]', '', text)

        # Normalize line breaks
        text = re.sub(r'\n{3,}', '\n\n', text)

        return text.strip()

    def _detect_sections(self, text: str) -> Dict[str, str]:
        """Detect and extract resume sections"""
        sections = {}
        lines = text.split('\n')

        current_section = "header"
        current_content = []

        for line in lines:
            line_stripped = line.strip()

            # Check if this line is a section header
            detected_section = None
            for section_name, pattern in self.SECTION_PATTERNS.items():
                if re.match(pattern, line_stripped):
                    detected_section = section_name
                    break

            if detected_section:
                # Save previous section
                if current_content:
                    sections[current_section] = "\n".join(current_content).strip()

                current_section = detected_section
                current_content = []
            else:
                current_content.append(line)

        # Save last section
        if current_content:
            sections[current_section] = "\n".join(current_content).strip()

        return sections

    def _extract_contact_info(self, text: str) -> Dict[str, Optional[str]]:
        """Extract contact information from resume"""
        contact = {
            "email": None,
            "phone": None,
            "linkedin": None,
            "github": None,
            "website": None,
            "location": None
        }

        # Email pattern
        email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
        if email_match:
            contact["email"] = email_match.group()

        # Phone pattern (various formats)
        phone_match = re.search(r'[\+]?[(]?[0-9]{1,3}[)]?[-\s\.]?[(]?[0-9]{1,3}[)]?[-\s\.]?[0-9]{3,4}[-\s\.]?[0-9]{3,4}', text)
        if phone_match:
            contact["phone"] = phone_match.group()

        # LinkedIn
        linkedin_match = re.search(r'linkedin\.com/in/[\w-]+', text, re.IGNORECASE)
        if linkedin_match:
            contact["linkedin"] = linkedin_match.group()

        # GitHub
        github_match = re.search(r'github\.com/[\w-]+', text, re.IGNORECASE)
        if github_match:
            contact["github"] = github_match.group()

        return contact


class ResumeSectionExtractor:
    """Helper class for more detailed section extraction"""

    @staticmethod
    def extract_experience_entries(experience_text: str) -> List[Dict]:
        """Extract individual experience entries"""
        entries = []

        # Pattern for common date formats
        date_pattern = r'(\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}\s*[-–]\s*(?:Present|Current|\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4})\b)'

        # Split by date patterns
        parts = re.split(date_pattern, experience_text)

        current_entry = {}
        for i, part in enumerate(parts):
            if re.match(date_pattern, part):
                if current_entry:
                    entries.append(current_entry)
                current_entry = {"date_range": part.strip()}
            elif current_entry:
                current_entry["content"] = part.strip()

        if current_entry and "content" in current_entry:
            entries.append(current_entry)

        return entries

    @staticmethod
    def extract_skills_list(skills_text: str) -> List[str]:
        """Extract individual skills from skills section"""
        # Common separators
        separators = [',', '|', '•', '·', '-', '\n']

        skills = [skills_text]
        for sep in separators:
            new_skills = []
            for skill in skills:
                new_skills.extend(skill.split(sep))
            skills = new_skills

        # Clean and filter
        skills = [s.strip() for s in skills if s.strip() and len(s.strip()) > 1]

        return skills
